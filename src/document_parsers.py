"""
Document parsers for different file formats (PDF, DOCX, Email)
"""
import io
import re
import uuid
import hashlib
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import logging

# PDF parsing
import PyPDF2

# DOCX parsing
from docx import Document as DocxDocument

# Email parsing
from email import message_from_string
from email.policy import default
import email_reply_parser

from .models import Document, DocumentType, Domain
from .config import config

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors"""
    pass


class BaseDocumentParser:
    """Base class for document parsers"""
    
    def __init__(self):
        self.supported_extensions = []
    
    def parse(self, content: Union[str, bytes], filename: str = "", metadata: Dict[str, Any] = None) -> Document:
        """Parse document content and return a Document object"""
        raise NotImplementedError
    
    def _generate_document_id(self, content: str, filename: str = "") -> str:
        """Generate unique document ID based on content hash"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        return f"{filename}_{content_hash[:8]}" if filename else content_hash[:16]
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\|\\\^\`\~]', '', text)
        return text.strip()
    
    def _extract_metadata(self, content: str, filename: str = "") -> Dict[str, Any]:
        """Extract basic metadata from content"""
        metadata = {
            'filename': filename,
            'word_count': len(content.split()),
            'char_count': len(content),
            'has_tables': 'table' in content.lower() or '|' in content,
            'has_lists': bool(re.search(r'^\s*[\*\-\d+]\s+', content, re.MULTILINE)),
        }
        return metadata


class PDFParser(BaseDocumentParser):
    """Parser for PDF documents"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def parse(self, content: Union[str, bytes], filename: str = "", metadata: Dict[str, Any] = None) -> Document:
        """Parse PDF content"""
        try:
            if isinstance(content, str):
                # If content is a file path
                with open(content, 'rb') as file:
                    pdf_content = file.read()
            else:
                pdf_content = content
            
            # Create PDF reader
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
            
            # Extract text from all pages
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content.strip():
                raise DocumentParsingError("No text content could be extracted from PDF")
            
            # Clean text
            cleaned_content = self._clean_text(text_content)
            
            # Extract metadata
            doc_metadata = self._extract_metadata(cleaned_content, filename)
            doc_metadata.update({
                'page_count': len(pdf_reader.pages),
                'pdf_metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
            })
            
            if metadata:
                doc_metadata.update(metadata)
            
            # Create document
            document = Document(
                id=self._generate_document_id(cleaned_content, filename),
                title=filename or "Untitled PDF",
                content=cleaned_content,
                doc_type=DocumentType.PDF,
                domain=self._infer_domain(cleaned_content),
                source_path=filename if Path(filename).exists() else None,
                metadata=doc_metadata
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise DocumentParsingError(f"Failed to parse PDF: {str(e)}")


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX documents"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def parse(self, content: Union[str, bytes], filename: str = "", metadata: Dict[str, Any] = None) -> Document:
        """Parse DOCX content"""
        try:
            if isinstance(content, str):
                # If content is a file path
                doc = DocxDocument(content)
            else:
                # If content is bytes
                doc = DocxDocument(io.BytesIO(content))
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_cells))
                tables_text.append("\n".join(table_rows))
            
            # Combine all text
            full_content = "\n\n".join(paragraphs)
            if tables_text:
                full_content += "\n\n--- TABLES ---\n\n" + "\n\n".join(tables_text)
            
            if not full_content.strip():
                raise DocumentParsingError("No text content could be extracted from DOCX")
            
            # Clean text
            cleaned_content = self._clean_text(full_content)
            
            # Extract metadata
            doc_metadata = self._extract_metadata(cleaned_content, filename)
            doc_metadata.update({
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'docx_properties': self._extract_docx_properties(doc)
            })
            
            if metadata:
                doc_metadata.update(metadata)
            
            # Create document
            document = Document(
                id=self._generate_document_id(cleaned_content, filename),
                title=filename or "Untitled DOCX",
                content=cleaned_content,
                doc_type=DocumentType.DOCX,
                domain=self._infer_domain(cleaned_content),
                source_path=filename if Path(filename).exists() else None,
                metadata=doc_metadata
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}")
    
    def _extract_docx_properties(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract document properties from DOCX"""
        properties = {}
        try:
            core_props = doc.core_properties
            properties.update({
                'author': core_props.author or '',
                'title': core_props.title or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            })
        except Exception as e:
            logger.warning(f"Could not extract DOCX properties: {e}")
        
        return properties


class EmailParser(BaseDocumentParser):
    """Parser for email messages"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.eml', '.msg', '.txt']
    
    def parse(self, content: Union[str, bytes], filename: str = "", metadata: Dict[str, Any] = None) -> Document:
        """Parse email content"""
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            # Parse email message
            if content.strip().startswith(('From:', 'To:', 'Subject:', 'Date:')):
                # Looks like email headers
                msg = message_from_string(content, policy=default)
            else:
                # Treat as plain text email
                msg = None
            
            # Extract email components
            if msg:
                subject = msg.get('Subject', 'No Subject')
                sender = msg.get('From', 'Unknown Sender')
                recipient = msg.get('To', 'Unknown Recipient')
                date = msg.get('Date', 'Unknown Date')
                
                # Extract body
                body = self._extract_email_body(msg)
            else:
                # Fallback for plain text
                subject = filename or "Email Message"
                sender = "Unknown"
                recipient = "Unknown"
                date = "Unknown"
                body = content
            
            # Parse reply and remove quoted text
            try:
                clean_body = email_reply_parser.EmailReplyParser.parse_reply(body)
            except:
                clean_body = body
            
            # Format email content
            email_content = f"""Subject: {subject}
From: {sender}
To: {recipient}
Date: {date}

{clean_body}"""
            
            if not email_content.strip():
                raise DocumentParsingError("No content could be extracted from email")
            
            # Clean text
            cleaned_content = self._clean_text(email_content)
            
            # Extract metadata
            doc_metadata = self._extract_metadata(cleaned_content, filename)
            doc_metadata.update({
                'email_subject': subject,
                'email_from': sender,
                'email_to': recipient,
                'email_date': date,
                'has_attachments': msg.get_content_disposition() == 'attachment' if msg else False
            })
            
            if metadata:
                doc_metadata.update(metadata)
            
            # Create document
            document = Document(
                id=self._generate_document_id(cleaned_content, filename),
                title=subject,
                content=cleaned_content,
                doc_type=DocumentType.EMAIL,
                domain=self._infer_domain(cleaned_content),
                source_path=filename if Path(filename).exists() else None,
                metadata=doc_metadata
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Error parsing email: {e}")
            raise DocumentParsingError(f"Failed to parse email: {str(e)}")
    
    def _extract_email_body(self, msg) -> str:
        """Extract body text from email message"""
        body = ""
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get_content_disposition())
                
                if content_type == "text/plain" and "attachment" not in content_disposition:
                    try:
                        body += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    except:
                        body += str(part.get_payload())
        else:
            try:
                body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            except:
                body = str(msg.get_payload())
        
        return body
    
    def _infer_domain(self, content: str) -> Domain:
        """Infer document domain based on content analysis"""
        content_lower = content.lower()
        
        # Insurance keywords
        insurance_keywords = ['policy', 'coverage', 'claim', 'premium', 'deductible', 'insurance', 'benefit']
        insurance_score = sum(1 for keyword in insurance_keywords if keyword in content_lower)
        
        # Legal keywords
        legal_keywords = ['contract', 'agreement', 'clause', 'legal', 'liability', 'breach', 'jurisdiction']
        legal_score = sum(1 for keyword in legal_keywords if keyword in content_lower)
        
        # HR keywords
        hr_keywords = ['employee', 'salary', 'benefit', 'leave', 'performance', 'evaluation', 'hr', 'human resources']
        hr_score = sum(1 for keyword in hr_keywords if keyword in content_lower)
        
        # Compliance keywords
        compliance_keywords = ['compliance', 'regulation', 'audit', 'risk', 'standard', 'procedure', 'requirement']
        compliance_score = sum(1 for keyword in compliance_keywords if keyword in content_lower)
        
        # Determine domain based on highest score
        scores = {
            Domain.INSURANCE: insurance_score,
            Domain.LEGAL: legal_score,
            Domain.HR: hr_score,
            Domain.COMPLIANCE: compliance_score
        }
        
        return max(scores, key=scores.get) if max(scores.values()) > 0 else Domain.LEGAL


class DocumentParserFactory:
    """Factory class for creating appropriate document parsers"""
    
    def __init__(self):
        self.parsers = {
            '.pdf': PDFParser(),
            '.docx': DOCXParser(),
            '.eml': EmailParser(),
            '.msg': EmailParser(),
            '.txt': EmailParser(),  # Can handle plain text emails
        }
    
    def get_parser(self, file_extension: str) -> BaseDocumentParser:
        """Get appropriate parser for file extension"""
        extension = file_extension.lower()
        if extension not in self.parsers:
            raise DocumentParsingError(f"No parser available for extension: {extension}")
        return self.parsers[extension]
    
    def parse_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Document:
        """Parse document from file path"""
        path = Path(file_path)
        if not path.exists():
            raise DocumentParsingError(f"File not found: {file_path}")
        
        parser = self.get_parser(path.suffix)
        return parser.parse(str(path), filename=path.name, metadata=metadata)
    
    def parse_content(self, content: Union[str, bytes], file_extension: str, 
                     filename: str = "", metadata: Dict[str, Any] = None) -> Document:
        """Parse document from content"""
        parser = self.get_parser(file_extension)
        return parser.parse(content, filename=filename, metadata=metadata)